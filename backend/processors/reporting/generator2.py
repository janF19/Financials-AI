from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg') # Set non-interactive backend before importing pyplot
import matplotlib.pyplot as plt
import io
from typing import Dict, Any, Optional, List
from openai import OpenAI
from backend.config.settings import settings # Assuming settings are configured
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class ReportGenerator2:

    @staticmethod
    def _get_openai_client() -> Optional[OpenAI]:
        try:
            if settings.OPENAI_API_KEY:
                client = OpenAI(api_key=settings.OPENAI_API_KEY)
                client.models.list() # Verify key
                return client
            else:
                logger.error("OpenAI API key not found in settings.")
                return None
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI client: {e}")
            return None

    @staticmethod
    def _add_heading(doc: Document, text: str, level: int):
        heading = doc.add_heading(text, level=level)
        if level == 0: # Main Title
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.runs[0]
            run.font.size = Pt(24)
            run.font.bold = True
        elif level == 1: # Major Sections
            run = heading.runs[0]
            run.font.size = Pt(16)
            run.font.bold = True
            # Add a bottom border to H1
            p = heading._p
            pPr = p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom_bdr = OxmlElement('w:bottom')
            bottom_bdr.set(qn('w:val'), 'single')
            bottom_bdr.set(qn('w:sz'), '6') # border size
            bottom_bdr.set(qn('w:space'), '1')
            bottom_bdr.set(qn('w:color'), 'auto')
            pbdr.append(bottom_bdr)
            pPr.append(pbdr)
        elif level == 2: # Subsections
            run = heading.runs[0]
            run.font.size = Pt(13)
            run.font.bold = True
        # Add some space after headings
        doc.add_paragraph()


    @staticmethod
    def _format_currency(value: Any, currency_symbol: str = "Kč", unit: str = " thousands") -> str:
        if value is None or value == 'N/A':
            return "N/A"
        try:
            num_value = float(value)
            return f"{num_value:,.0f} {currency_symbol}{unit}"
        except (ValueError, TypeError):
            return str(value) # Return as is if not a number

    @staticmethod
    def _add_table(doc: Document, data_dict: Dict[str, Any], headers: List[str]):
        if not data_dict or all(v is None or v == 'N/A' for v in data_dict.values()):
            doc.add_paragraph("Data not available for this table.")
            return

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        for key, value in data_dict.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            # Assuming value is for the second column, format if numeric
            row_cells[1].text = ReportGenerator2._format_currency(value, unit="") if len(headers) == 2 else str(value)
        doc.add_paragraph()


    @staticmethod
    def generate(data: Dict[str, Any]) -> Document:
        doc = Document()
        client = ReportGenerator2._get_openai_client()

        # --- Overall Document Structure ---
        # Part 1: Executive Summary & Key Findings (Target ~2 pages)
        # Part 2: Detailed Analysis & Supporting Data

        # --- Metadata & Title ---
        ReportGenerator2._add_heading(doc, 'Company Valuation & Analysis Report', level=0)
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(f"Report Generated: {timestamp}\n").italic = True
        
        # Safely access data
        company_info = data.get('information', {})
        financial_data = data # This is the top-level dict from extractor
        valuation_results = data.get('result_valuation', {})
        document_analysis = data.get('document_analysis', {})
        income_statement_data = data.get('income_statement', {})
        balance_sheet_data = data.get('balance_sheet', {})
        cash_flow_data = data.get('cash_flow', {})

        # =============================================
        # PART 1: EXECUTIVE SUMMARY & KEY FINDINGS
        # =============================================
        ReportGenerator2._add_heading(doc, 'Part 1: Executive Summary & Key Findings', level=1)

        # 1.1 Company Snapshot
        ReportGenerator2._add_heading(doc, 'Company Snapshot', level=2)
        cs_para = doc.add_paragraph()
        cs_para.add_run(f"Company Name: ").bold = True
        cs_para.add_run(f"{company_info.get('company_name', 'N/A')}\n")
        cs_para.add_run(f"Industry: ").bold = True
        cs_para.add_run(f"{company_info.get('industry', 'N/A')}\n")
        cs_para.add_run(f"Accounting Period: ").bold = True
        cs_para.add_run(f"{company_info.get('accounting_period', 'N/A')}\n")
        cs_para.add_run(f"Headquarters: ").bold = True
        cs_para.add_run(f"{company_info.get('headquarters', 'N/A')}\n")
        
        company_summary_ai = company_info.get('analytical_summary', 'N/A')
        if company_summary_ai and company_summary_ai != 'N/A':
            cs_para.add_run(f"Company Profile Summary (AI Generated):\n").bold = True
            cs_para.add_run(company_summary_ai)
        else:
            # Fallback if AI summary is not available
            desc = ReportGenerator2._generate_text_summary(
                client,
                f"Provide a brief 2-sentence professional description for {company_info.get('company_name', 'this company')} operating in the {company_info.get('industry', 'relevant')} industry, with main activities: {', '.join(company_info.get('main_activities', ['not specified']))}.",
                "Company Description"
            )
            cs_para.add_run(f"Company Profile Summary (AI Generated):\n").bold = True
            cs_para.add_run(desc)
        doc.add_paragraph()

        # 1.2 Key Valuation Results
        ReportGenerator2._add_heading(doc, 'Key Valuation Results', level=2)
        kv_para = doc.add_paragraph()
        ebit_val = valuation_results.get('Enterprise_Value_based_on_EBIT_Kč_thousands')
        ebitda_val = valuation_results.get('Enterprise_Value_based_on_EBITDA_Kč_thousands')
        
        kv_para.add_run("Enterprise Value (EV) Estimates (Adjusted for 2025):\n").bold = True
        kv_para.add_run(f"  • Based on EV/EBIT ({valuation_results.get('EV_EBIT_Multiple', 'N/A')}x): ")
        kv_para.add_run(ReportGenerator2._format_currency(ebit_val) + "\n")
        kv_para.add_run(f"  • Based on EV/EBITDA ({valuation_results.get('EV_EBITDA_Multiple', 'N/A')}x): ")
        kv_para.add_run(ReportGenerator2._format_currency(ebitda_val) + "\n")
        
        if isinstance(valuation_results.get('EBIT_adjusted_for_2025'), (int, float)) and \
           isinstance(valuation_results.get('EBITDA_adjusted_for_2025'), (int, float)):
            ReportGenerator2._add_valuation_graph(doc, valuation_results, "EBIT_adjusted_for_2025", "EBITDA_adjusted_for_2025", "Adjusted EBIT vs EBITDA (for 2025)")
        else:
            doc.add_paragraph("Adjusted EBIT/EBITDA graph data not available.")
        doc.add_paragraph()

        # 1.3 Overall Financial Health & Performance (AI Generated)
        ReportGenerator2._add_heading(doc, 'Overall Financial Health & Performance', level=2)
        health_prompt_context = f"""
        Company: {company_info.get('company_name', 'N/A')}
        Industry: {company_info.get('industry', 'N/A')}
        Key Financials (Current Period):
        - Revenue (Products & Services): {ReportGenerator2._format_currency(income_statement_data.get('revenue_from_products_and_services_current'))}
        - Operating Profit (EBIT): {ReportGenerator2._format_currency(income_statement_data.get('operating_profit_current'))}
        - Net Operating Cash Flow: {ReportGenerator2._format_currency(cash_flow_data.get('net_operating_cash_flow_current'))}
        - Total Assets: {ReportGenerator2._format_currency(balance_sheet_data.get('total_assets_current'))}
        - Total Equity: {ReportGenerator2._format_currency(balance_sheet_data.get('equity_current'))}
        
        Management Discussion Summary: {document_analysis.get('management_discussion_summary', 'N/A')}
        Income Statement Analysis: {income_statement_data.get('analytical_summary', 'N/A')}
        Balance Sheet Analysis: {balance_sheet_data.get('analytical_summary', 'N/A')}
        Cash Flow Analysis: {cash_flow_data.get('analytical_summary', 'N/A')}
        """
        health_analysis = ReportGenerator2._generate_text_summary(
            client,
            f"Based on the following data, provide a concise (3-5 sentences) assessment of the company's overall financial health and performance. Highlight key strengths and weaknesses. \nContext:\n{health_prompt_context}",
            "Financial Health Assessment"
        )
        doc.add_paragraph(health_analysis)
        doc.add_paragraph()

        # 1.4 Key Strategic Insights & Outlook (from document_analysis)
        ReportGenerator2._add_heading(doc, 'Key Strategic Insights & Outlook', level=2)
        doc.add_paragraph().add_run("Management Discussion Summary:").bold = True
        doc.add_paragraph(document_analysis.get('management_discussion_summary', 'N/A'))
        doc.add_paragraph().add_run("Significant Events/Achievements:").bold = True
        doc.add_paragraph(document_analysis.get('significant_events_achievements', 'N/A'))
        doc.add_paragraph().add_run("Key Risks & Uncertainties:").bold = True
        doc.add_paragraph(document_analysis.get('key_risks_and_uncertainties', 'N/A'))
        doc.add_paragraph().add_run("Future Outlook & Strategy:").bold = True
        doc.add_paragraph(document_analysis.get('future_outlook_and_strategy', 'N/A'))
        doc.add_paragraph()

        # 1.5 Conclusion (AI Generated)
        ReportGenerator2._add_heading(doc, 'Valuation Conclusion', level=2)
        conclusion_prompt_context = f"""
        Company: {company_info.get('company_name', 'N/A')}
        Valuation based on EV/EBIT: {ReportGenerator2._format_currency(ebit_val)}
        Valuation based on EV/EBITDA: {ReportGenerator2._format_currency(ebitda_val)}
        Overall Financial Health Assessment: {health_analysis}
        Key Strategic Insights: {document_analysis.get('management_discussion_summary', 'N/A')}
        """
        conclusion = ReportGenerator2._generate_text_summary(
            client,
            f"Based on the valuation results and overall analysis provided below, write a brief (2-4 sentences) concluding statement for this valuation report. \nContext:\n{conclusion_prompt_context}",
            "Report Conclusion"
        )
        doc.add_paragraph(conclusion)

        # End of Part 1 - Add Page Break
        doc.add_page_break()

        # =============================================
        # PART 2: DETAILED ANALYSIS & SUPPORTING DATA
        # =============================================
        ReportGenerator2._add_heading(doc, 'Part 2: Detailed Analysis & Supporting Data', level=1)

        # 2.1 Detailed Company Overview
        ReportGenerator2._add_heading(doc, 'Detailed Company Overview', level=2)
        dco_para = doc.add_paragraph()
        dco_para.add_run("Company Name: ").bold = True
        dco_para.add_run(f"{company_info.get('company_name', 'N/A')}\n")
        # ... Add more fields from company_info as needed ...
        dco_para.add_run("IČ (ID): ").bold = True
        dco_para.add_run(f"{company_info.get('IC', 'N/A')}\n")
        dco_para.add_run("Legal Form: ").bold = True
        dco_para.add_run(f"{company_info.get('legal_form', 'N/A')}\n")
        dco_para.add_run("Registered Capital: ").bold = True
        dco_para.add_run(f"{company_info.get('registered_capital', 'N/A')}\n")
        dco_para.add_run("Established: ").bold = True
        dco_para.add_run(f"{company_info.get('established', 'N/A')}\n")
        dco_para.add_run("Employee Count: ").bold = True
        dco_para.add_run(f"{company_info.get('employee_count', 'N/A')}\n")
        dco_para.add_run("Main Activities: ").bold = True
        dco_para.add_run(f"{', '.join(company_info.get('main_activities', ['N/A']))}\n")
        dco_para.add_run("Recent News/Developments (from report context): ").bold = True
        dco_para.add_run(f"{company_info.get('news', 'N/A')}\n")
        doc.add_paragraph()

        # 2.2 Detailed Financial Statement Analysis
        # Income Statement
        ReportGenerator2._add_heading(doc, 'Income Statement Analysis', level=2)
        is_table_data = {
            "Revenue (Products & Services)": income_statement_data.get('revenue_from_products_and_services_current'),
            "Revenue (Goods)": income_statement_data.get('revenue_from_goods_current'),
            "Production Consumption": income_statement_data.get('production_consumption_current'),
            "Personnel Costs": income_statement_data.get('personnel_costs_current'),
            "Depreciation & Amortization": income_statement_data.get('depreciation_current'),
            "Operating Profit (EBIT)": income_statement_data.get('operating_profit_current'),
        }
        ReportGenerator2._add_table(doc, is_table_data, ["Metric (Current Period)", "Value (Kč thousands)"])
        doc.add_paragraph().add_run("AI Generated Analysis:").bold = True
        doc.add_paragraph(income_statement_data.get('analytical_summary', 'No specific AI analysis available for the income statement.'))
        doc.add_paragraph()

        # Balance Sheet
        ReportGenerator2._add_heading(doc, 'Balance Sheet Analysis', level=2)
        # For balance sheet, we might want current and previous columns
        bs_table = doc.add_table(rows=1, cols=3)
        bs_table.style = 'Table Grid'
        bs_table.cell(0,0).text = "Metric"
        bs_table.cell(0,1).text = "Current Period (Kč thousands)"
        bs_table.cell(0,2).text = "Previous Period (Kč thousands)"
        bs_metrics = [
            ("Total Assets", 'total_assets_current', 'total_assets_previous'),
            ("Tangible Assets", 'tangible_assets_current', 'tangible_assets_previous'),
            ("Current Assets", 'current_assets_current', 'current_assets_previous'),
            ("Total Equity", 'equity_current', 'equity_previous'),
            ("Total Liabilities", 'liabilities_current', 'liabilities_previous'),
        ]
        for name, cur_key, prev_key in bs_metrics:
            row = bs_table.add_row().cells
            row[0].text = name
            row[1].text = ReportGenerator2._format_currency(balance_sheet_data.get(cur_key), unit="")
            row[2].text = ReportGenerator2._format_currency(balance_sheet_data.get(prev_key), unit="")
        doc.add_paragraph()
        doc.add_paragraph().add_run("AI Generated Analysis:").bold = True
        doc.add_paragraph(balance_sheet_data.get('analytical_summary', 'No specific AI analysis available for the balance sheet.'))
        doc.add_paragraph()

        # Cash Flow Statement
        ReportGenerator2._add_heading(doc, 'Cash Flow Statement Analysis', level=2)
        cf_table = doc.add_table(rows=1, cols=3)
        cf_table.style = 'Table Grid'
        cf_table.cell(0,0).text = "Metric"
        cf_table.cell(0,1).text = "Current Period (Kč thousands)"
        cf_table.cell(0,2).text = "Previous Period (Kč thousands)"
        cf_metrics = [
            ("Initial Cash Balance", 'initial_cash_balance_current', 'initial_cash_balance_previous'),
            ("Net Operating Cash Flow", 'net_operating_cash_flow_current', 'net_operating_cash_flow_previous'),
            ("CAPEX (Investment)", 'capex_current', 'capex_previous'),
            ("Proceeds from Sale of Fixed Assets", 'proceeds_from_sale_of_fixed_assets_current', 'proceeds_from_sale_of_fixed_assets_previous'),
        ]
        for name, cur_key, prev_key in cf_metrics:
            row = cf_table.add_row().cells
            row[0].text = name
            row[1].text = ReportGenerator2._format_currency(cash_flow_data.get(cur_key), unit="")
            row[2].text = ReportGenerator2._format_currency(cash_flow_data.get(prev_key), unit="")
        doc.add_paragraph()
        doc.add_paragraph().add_run("AI Generated Analysis:").bold = True
        doc.add_paragraph(cash_flow_data.get('analytical_summary', 'No specific AI analysis available for the cash flow statement.'))
        doc.add_paragraph()

        # 2.3 Detailed Valuation Metrics
        ReportGenerator2._add_heading(doc, 'Detailed Valuation Metrics', level=2)
        val_details_para = doc.add_paragraph()
        val_details_para.add_run("Original Financials (for valuation basis):\n").bold = True
        ebit_orig = valuation_results.get('EBIT_original', {})
        ebitda_orig = valuation_results.get('EBITDA_original', {})
        val_details_para.add_run(f"  • EBIT ({ebit_orig.get('year', 'N/A')}): {ReportGenerator2._format_currency(ebit_orig.get('value'))}\n")
        val_details_para.add_run(f"  • EBITDA ({ebitda_orig.get('year', 'N/A')}): {ReportGenerator2._format_currency(ebitda_orig.get('value'))}\n")
        
        val_details_para.add_run("Adjusted Financials (for 2025 multiples):\n").bold = True
        val_details_para.add_run(f"  • Adjusted EBIT (for 2025): {ReportGenerator2._format_currency(valuation_results.get('EBIT_adjusted_for_2025'))}\n")
        val_details_para.add_run(f"  • Adjusted EBITDA (for 2025): {ReportGenerator2._format_currency(valuation_results.get('EBITDA_adjusted_for_2025'))}\n")

        val_details_para.add_run("Applied Multiples:\n").bold = True
        val_details_para.add_run(f"  • EV/EBIT Multiple: {valuation_results.get('EV_EBIT_Multiple', 'N/A')}x\n")
        val_details_para.add_run(f"  • EV/EBITDA Multiple: {valuation_results.get('EV_EBITDA_Multiple', 'N/A')}x\n")
        doc.add_paragraph()

        # Disclaimer
        ReportGenerator2._add_heading(doc, 'Disclaimer', level=2)
        doc.add_paragraph(
            "This report is generated based on automated analysis of financial documents and publicly available data. "
            "The information and valuation provided are for informational purposes only and should not be considered as financial advice. "
            "Users should conduct their own due diligence before making any investment decisions."
        ).italic = True

        return doc

    @staticmethod
    def _generate_text_summary(client: Optional[OpenAI], prompt: str, summary_type: str, model: str = "gpt-4o") -> str:
        if not client:
            logger.warning(f"OpenAI client not available. Cannot generate {summary_type}.")
            return f"AI-generated {summary_type} could not be produced as the AI service is unavailable."
        try:
            logger.info(f"Generating AI summary for: {summary_type}")
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a concise financial analyst assistant. Provide clear and brief summaries based on the context. Focus on the key takeaways."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3, # Slightly creative but still factual
                max_tokens=250 # Adjust as needed
            )
            summary = response.choices[0].message.content.strip()
            logger.info(f"Successfully generated AI summary for: {summary_type}")
            return summary
        except Exception as e:
            logger.error(f"Error generating AI summary for {summary_type}: {e}")
            return f"An error occurred while generating the {summary_type}."

    @staticmethod
    def _add_valuation_graph(doc: Document, valuation_data: Dict, ebit_key: str, ebitda_key: str, title: str):
        try:
            ebit_value = valuation_data.get(ebit_key)
            ebitda_value = valuation_data.get(ebitda_key)

            if not isinstance(ebit_value, (int, float)) or not isinstance(ebitda_value, (int, float)):
                doc.add_paragraph(f"Insufficient data for '{title}' graph.")
                logger.warning(f"Graph generation skipped for '{title}' due to missing/invalid data: EBIT={ebit_value}, EBITDA={ebitda_value}")
                return

            plt.figure(figsize=(6, 3.5)) # Adjusted size
            metrics = ['EBIT', 'EBITDA']
            values = [ebit_value, ebitda_value]
            
            bars = plt.bar(metrics, values, color=['#4A86E8', '#76A7F2']) # Example colors
            plt.title(title, fontsize=10)
            plt.ylabel('Value (Kč thousands)', fontsize=8)
            plt.xticks(fontsize=8)
            plt.yticks(fontsize=8)
            plt.grid(axis='y', linestyle='--', alpha=0.7)

            # Add labels on bars
            for bar in bars:
                yval = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2.0, yval + (max(values)*0.02), f'{yval:,.0f}', ha='center', va='bottom', fontsize=7)

            img_stream = io.BytesIO()
            plt.savefig(img_stream, format='png', bbox_inches='tight', dpi=150)
            img_stream.seek(0)
            
            doc.add_picture(img_stream, width=Inches(5.0)) # Adjusted width
            plt.close() # Close the plot to free memory
            logger.info(f"Successfully added graph: {title}")
        except Exception as e:
            doc.add_paragraph(f"Error generating graph '{title}': {str(e)}")
            logger.error(f"Error generating graph '{title}': {e}", exc_info=True)

if __name__ == '__main__':
    # Dummy data for testing ReportGenerator2
    test_data = {
        "information": {
            "company_name": "TestCorp AG", "industry": "Software Development",
            "accounting_period": "2023 Full Year", "headquarters": "Prague, Czech Republic",
            "IC": "12345678", "legal_form": "a.s.", "registered_capital": "1,000,000 Kč",
            "established": "01.01.2000", "employee_count": 150,
            "main_activities": ["Custom software solutions", "AI research"],
            "news": "Launched new AI platform 'Cognito' in Q3 2023.",
            "analytical_summary": "TestCorp AG is a leading software firm specializing in AI, showing strong growth with its new Cognito platform."
        },
        "income_statement": {
            "revenue_from_products_and_services_current": 50000, "operating_profit_current": 10000,
            "depreciation_current": 2000, "personnel_costs_current": 25000,
            "analytical_summary": "Revenue driven by Cognito platform. Personnel costs are significant but stable. Profitability is healthy."
        },
        "balance_sheet": {
            "total_assets_current": 100000, "total_assets_previous": 80000,
            "equity_current": 60000, "equity_previous": 55000,
            "analytical_summary": "Assets grew due to increased cash reserves. Equity position strengthened."
        },
        "cash_flow": {
            "net_operating_cash_flow_current": 12000, "net_operating_cash_flow_previous": 9000,
            "capex_current": -3000, "capex_previous": -2000,
            "analytical_summary": "Strong operating cash flow supports investments and debt repayment."
        },
        "document_analysis": {
            "management_discussion_summary": "Management is optimistic about future growth driven by AI. Focus on expanding market share for Cognito.",
            "significant_events_achievements": "Successful launch of Cognito platform. Secured two major enterprise clients.",
            "key_risks_and_uncertainties": "Competition in AI space, talent retention.",
            "future_outlook_and_strategy": "Plan to expand into new international markets in 2024. Continued R&D investment in AI."
        },
        "result_valuation": {
            "EBIT_original": {"value": 10000, "year": 2023},
            "EBITDA_original": {"value": 12000, "year": 2023},
            "EBIT_adjusted_for_2025": 11500, # Dummy adjusted values
            "EBITDA_adjusted_for_2025": 13800,
            "EV_EBIT_Multiple": 10.5, "EV_EBITDA_Multiple": 8.5,
            "Enterprise_Value_based_on_EBIT_Kč_thousands": 120750, # 11500 * 10.5
            "Enterprise_Value_based_on_EBITDA_Kč_thousands": 117300 # 13800 * 8.5
        }
    }
    # Ensure OPENAI_API_KEY is set in your environment or backend.config.settings for AI summaries to work
    if not settings.OPENAI_API_KEY:
        logger.warning("OPENAI_API_KEY not set. AI-generated summaries in the report will be placeholders.")

    report_doc = ReportGenerator2.generate(test_data)
    output_filename = "Test_Company_Report_V2.docx"
    report_doc.save(output_filename)
    logger.info(f"Test report V2 saved as {output_filename}")
    print(f"Test report V2 saved as {output_filename}") 